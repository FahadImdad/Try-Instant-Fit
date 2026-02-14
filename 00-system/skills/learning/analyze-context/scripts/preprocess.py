#!/usr/bin/env python3
"""
Pre-process uploaded files for SubAgent analysis.

Converts PDFs and DOCX to text, chunks large files.
Run from Nexus root directory.

Usage:
    uv run 00-system/skills/learning/analyze-context/scripts/preprocess.py

Output:
    - Extracted text files in 01-memory/input/_analysis/extracted/
    - Prints JSON list of processable file paths
"""

import json
import sys
from pathlib import Path

# Configuration
INPUT_DIR = Path("01-memory/input")
EXTRACTED_DIR = INPUT_DIR / "_analysis" / "extracted"
CHUNK_SIZE_CHARS = 50000  # ~50KB, ~12,000 words per chunk

# Track statistics
stats = {
    "uploaded": 0,
    "pdfs_extracted": 0,
    "docx_extracted": 0,
    "chunked": 0,
    "skipped": 0,
    "processable": 0
}


def main():
    """Main entry point."""
    EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)

    # Get all files (not directories, not _analysis folder)
    uploaded_files = [
        f for f in INPUT_DIR.iterdir()
        if f.is_file() and not f.name.startswith('_')
    ]

    stats["uploaded"] = len(uploaded_files)

    if not uploaded_files:
        print("No files found in 01-memory/input/")
        print("Please upload files first.")
        sys.exit(1)

    processable_files = []

    for file_path in uploaded_files:
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            extracted = extract_pdf(file_path)
            if extracted:
                processable_files.extend(extracted)
                stats["pdfs_extracted"] += 1
            else:
                stats["skipped"] += 1

        elif suffix == '.docx':
            extracted = extract_docx(file_path)
            if extracted:
                processable_files.extend(extracted)
                stats["docx_extracted"] += 1
            else:
                stats["skipped"] += 1

        elif suffix in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            # Images: keep original (SubAgent uses multimodal)
            processable_files.append(str(file_path))

        else:
            # Text-based files: check size
            try:
                text = file_path.read_text(encoding='utf-8', errors='ignore')
                if len(text) > CHUNK_SIZE_CHARS:
                    chunks = chunk_text(file_path.stem, text)
                    processable_files.extend(chunks)
                    stats["chunked"] += 1
                else:
                    processable_files.append(str(file_path))
            except Exception as e:
                print(f"Warning: Could not read {file_path}: {e}", file=sys.stderr)
                stats["skipped"] += 1

    stats["processable"] = len(processable_files)

    # Print summary
    print("\n" + "=" * 50)
    print("PRE-PROCESSING COMPLETE")
    print("=" * 50)
    print(f"Uploaded files:     {stats['uploaded']}")
    print(f"PDFs extracted:     {stats['pdfs_extracted']}")
    print(f"DOCX extracted:     {stats['docx_extracted']}")
    print(f"Large files chunked:{stats['chunked']}")
    print(f"Skipped (errors):   {stats['skipped']}")
    print(f"Processable items:  {stats['processable']}")
    print("=" * 50 + "\n")

    # Output processable files as JSON (for AI to parse)
    print("PROCESSABLE_FILES_JSON:")
    print(json.dumps(processable_files, indent=2))

    return processable_files


def extract_pdf(file_path: Path) -> list:
    """Extract text from PDF, return list of output file paths."""
    try:
        import pdfplumber
    except ImportError:
        print(f"Warning: pdfplumber not installed. Run: pip install pdfplumber", file=sys.stderr)
        return []

    try:
        pdf = pdfplumber.open(file_path)
        full_text = ""

        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            full_text += f"\n--- Page {i+1} ---\n{page_text}\n"

        pdf.close()

        if not full_text.strip():
            print(f"Warning: No text extracted from {file_path} (may be scanned/image-based)", file=sys.stderr)
            return []

        # Check if chunking needed
        if len(full_text) > CHUNK_SIZE_CHARS:
            return chunk_text(file_path.stem, full_text)
        else:
            output_path = EXTRACTED_DIR / f"{file_path.stem}.txt"
            output_path.write_text(full_text, encoding='utf-8')
            return [str(output_path)]

    except Exception as e:
        print(f"Warning: PDF extraction failed for {file_path}: {e}", file=sys.stderr)
        return []


def extract_docx(file_path: Path) -> list:
    """Extract text from DOCX, return list of output file paths."""
    try:
        from docx import Document
    except ImportError:
        print(f"Warning: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        return []

    try:
        doc = Document(file_path)

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            print(f"Warning: No text extracted from {file_path}", file=sys.stderr)
            return []

        # Check if chunking needed
        if len(full_text) > CHUNK_SIZE_CHARS:
            return chunk_text(file_path.stem, full_text)
        else:
            output_path = EXTRACTED_DIR / f"{file_path.stem}.txt"
            output_path.write_text(full_text, encoding='utf-8')
            return [str(output_path)]

    except Exception as e:
        print(f"Warning: DOCX extraction failed for {file_path}: {e}", file=sys.stderr)
        return []


def chunk_text(base_name: str, text: str) -> list:
    """Split large text into chunks, return list of chunk file paths."""
    chunks = []
    chunk_num = 1

    # Split by double newlines (paragraphs) to avoid cutting mid-sentence
    paragraphs = text.split('\n\n')
    current_chunk = ""

    for para in paragraphs:
        # If adding this paragraph would exceed limit, save current chunk
        if current_chunk and len(current_chunk) + len(para) > CHUNK_SIZE_CHARS:
            chunk_path = EXTRACTED_DIR / f"{base_name}_chunk{chunk_num:02d}.txt"
            chunk_path.write_text(current_chunk, encoding='utf-8')
            chunks.append(str(chunk_path))
            chunk_num += 1
            current_chunk = para
        else:
            current_chunk = current_chunk + "\n\n" + para if current_chunk else para

    # Save final chunk
    if current_chunk.strip():
        chunk_path = EXTRACTED_DIR / f"{base_name}_chunk{chunk_num:02d}.txt"
        chunk_path.write_text(current_chunk, encoding='utf-8')
        chunks.append(str(chunk_path))

    print(f"Split '{base_name}' into {len(chunks)} chunks")
    return chunks


if __name__ == "__main__":
    main()
